import io

from docx import Document as DocxDocument

from app.services.extraction.base import ExtractedPage, ExtractionResult


def extract_docx(content: bytes) -> ExtractionResult:
    doc = DocxDocument(io.BytesIO(content))

    parts: list[str] = []

    for section in doc.sections:
        header = section.header
        if header is not None:
            for p in header.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    for section in doc.sections:
        footer = section.footer
        if footer is not None:
            for p in footer.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

    text = "\n".join(parts)
    # python-docx has no reliable page boundary API; DOCX is treated as a single logical page.
    page = ExtractedPage(page_number=1, text=text, ocr_used=False)
    return ExtractionResult(pages=[page])
