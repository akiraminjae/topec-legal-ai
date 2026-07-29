"""Draft response document generation from a case's integrated analysis (§22).

Deliberately minimal compared to the full spec: takes one free-text
`instructions` field rather than the full parameterized form (대응 대상
문서/인정 여부/문체/강도 등 개별 입력) — the AI is given the case analysis
summary plus whatever the user types in `instructions` and produces a single
draft. Structured per-field drafting input is a documented follow-up.
"""
from datetime import datetime

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.legal_case import CaseAnalysisSummary, LegalCase

DRAFT_DISCLAIMER = (
    "본 문서는 AI가 사건 통합분석 결과를 바탕으로 생성한 초안입니다.\n"
    "실제 제출 전 담당 변호사·법무담당자의 검토와 수정이 반드시 필요하며, 사실관계 확인 없이 그대로 제출해서는 안 됩니다."
)


def _write_header(doc: DocxDocument, subtitle: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TOPEC")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.bold = True
    sub_run.font.size = Pt(14)

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.add_run("AI 생성 초안 — 검토 전 사용 금지").italic = True

    doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')} (Asia/Seoul)")
    doc.add_paragraph("")


def build_case_response_draft(
    case: LegalCase, summary: CaseAnalysisSummary, report_type: str, instructions: str | None = None
) -> bytes:
    doc = DocxDocument()
    subtitle = "준비서면 초안" if report_type == "PREPARATORY_BRIEF_DRAFT" else "사건 경영진 보고 요약"
    _write_header(doc, subtitle)

    doc.add_heading("1. 사건 개요", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in (
        ("사건명", case.case_name),
        ("사건번호", case.case_number or "-"),
        ("법원", case.court_name or "-"),
        ("TOPEC의 지위", case.topec_position or "-"),
        ("상대방", case.opponent_name or "-"),
        ("청구금액", f"{case.claim_amount:,.0f} {case.currency}" if case.claim_amount is not None else "-"),
    ):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_heading("2. 사건 경과", level=1)
    doc.add_paragraph(summary.case_overview)

    doc.add_heading("3. 상대방 주장", level=1)
    doc.add_paragraph(summary.opponent_arguments_summary)

    doc.add_heading("4. TOPEC 입장", level=1)
    doc.add_paragraph(summary.topec_position_summary)

    doc.add_heading("5. 핵심 쟁점", level=1)
    doc.add_paragraph(summary.key_issues_summary)

    doc.add_heading("6. 누락 및 미대응사항", level=1)
    doc.add_paragraph(summary.missing_or_unaddressed)

    doc.add_heading("7. 종합 대응방향", level=1)
    doc.add_paragraph(summary.recommended_response_direction)

    if instructions:
        doc.add_heading("8. 담당자 지정 작성지침", level=1)
        doc.add_paragraph(instructions)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(DRAFT_DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)

    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
