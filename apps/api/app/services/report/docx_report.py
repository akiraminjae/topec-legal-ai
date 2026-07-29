"""DOCX report generation using python-docx.

No official TOPEC CI/logo file exists in this repository, so a plain text header
is used instead of a fabricated logo, per the "never invent branding assets"
principle. Swap `_write_header` for an image-based letterhead once an official
CI file is provided (see docs/DEPLOYMENT.md).
"""
import io
from datetime import datetime

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.document import Document
from app.models.enums import (
    CONTRACT_TYPE_LABELS_KO,
    LITIGATION_DOCUMENT_TYPE_LABELS_KO,
    REVISION_LEVEL_LABELS_KO,
    RISK_LEVEL_LABELS_KO,
    TOPEC_LITIGATION_POSITION_LABELS_KO,
    TOPEC_POSITION_LABELS_KO,
    ContractType,
    LitigationDocumentType,
    RevisionLevel,
    RiskLevel,
    TopecLitigationPosition,
    TopecPosition,
)

DISCLAIMER = (
    "본 결과는 AI를 활용한 1차 계약·법률 검토 지원자료입니다.\n"
    "사실관계, 계약상 지위 및 적용 법령에 따라 판단이 달라질 수 있습니다.\n"
    "중요 계약, 분쟁 가능 계약 또는 고위험 조항이 있는 경우 법무담당자나 외부 법률전문가의 확인을 거쳐야 합니다."
)


def _write_header(doc: DocxDocument, subtitle: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TOPEC")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.bold = True
    sub_run.font.size = Pt(14)

    mock_notice = doc.add_paragraph()
    mock_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mock_notice.add_run("AI 1차 검토 결과").italic = True

    doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')} (Asia/Seoul)")
    doc.add_paragraph("")


def _write_footer_disclaimer(doc: DocxDocument) -> None:
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)


def build_review_report(document: Document, summary, findings: list, revisions: list) -> bytes:
    is_litigation = document.document_category == "LITIGATION"
    doc = DocxDocument()
    _write_header(doc, "AI 소송·분쟁 검토 의견서" if is_litigation else "AI 계약·법률검토 의견서")

    doc.add_heading("1. 사건 개요" if is_litigation else "1. 계약 개요", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    if is_litigation:
        rows = [
            ("문서명", document.title),
            (
                "문서유형",
                LITIGATION_DOCUMENT_TYPE_LABELS_KO.get(
                    LitigationDocumentType(document.litigation_document_type), document.litigation_document_type
                ) if document.litigation_document_type else "-",
            ),
            (
                "TOPEC의 소송상 지위",
                TOPEC_LITIGATION_POSITION_LABELS_KO.get(
                    TopecLitigationPosition(document.topec_litigation_position), document.topec_litigation_position
                ) if document.topec_litigation_position else "-",
            ),
            ("사건번호", document.case_number or "-"),
            ("법원", document.court or "-"),
            ("상대방", document.counterparty_name or "-"),
            ("보안등급", document.security_level),
            ("전체 위험등급", RISK_LEVEL_LABELS_KO.get(RiskLevel(document.overall_risk_level), document.overall_risk_level or "-")
             if document.overall_risk_level else "-"),
        ]
    else:
        rows = [
            ("계약명", document.title),
            ("계약유형", CONTRACT_TYPE_LABELS_KO.get(ContractType(document.contract_type), document.contract_type)),
            ("TOPEC 계약상 지위", TOPEC_POSITION_LABELS_KO.get(TopecPosition(document.topec_position), document.topec_position)),
            ("상대방", document.counterparty_name or "-"),
            ("계약금액", f"{document.contract_amount:,.0f} {document.contract_currency}" if document.contract_amount else "-"),
            ("계약기간", f"{document.contract_start_date or '-'} ~ {document.contract_end_date or '-'}"),
            ("보안등급", document.security_level),
            ("전체 위험등급", RISK_LEVEL_LABELS_KO.get(RiskLevel(document.overall_risk_level), document.overall_risk_level or "-")
             if document.overall_risk_level else "-"),
        ]

    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)

    if summary:
        doc.add_heading("2. 사건 요약" if is_litigation else "2. 업무범위 요약", level=1)
        doc.add_paragraph(summary.scope_summary or "-")
        doc.add_heading("3. 핵심 쟁점 요약" if is_litigation else "3. 주요 위험 요약", level=1)
        doc.add_paragraph(summary.top_risks_summary or "-")

    doc.add_heading("4. 쟁점별 분석 및 대응방향" if is_litigation else "4. 조항별 위험분석", level=1)
    if not findings:
        doc.add_paragraph("탐지된 쟁점이 없습니다." if is_litigation else "탐지된 위험사항이 없습니다.")
    for i, f in enumerate(findings, start=1):
        doc.add_heading(f"{i}. {f.title} [{RISK_LEVEL_LABELS_KO.get(RiskLevel(f.risk_level), f.risk_level)}]", level=2)
        doc.add_paragraph(f"{'주장의 근거' if is_litigation else '위험 사유'}: {f.reason}")
        doc.add_paragraph(f"TOPEC에 미치는 영향: {f.impact_on_topec}")
        doc.add_paragraph(f"{'TOPEC 측 대응논리' if is_litigation else '권고 대응'}: {f.recommended_action}")
        doc.add_paragraph(f"AI 신뢰도: {f.confidence}% | 법무검토 필요: {'예' if f.legal_review_required else '아니오'}")

    if not is_litigation:
        doc.add_heading("5. 수정 전·후 비교", level=1)
        if not revisions:
            doc.add_paragraph("생성된 수정안이 없습니다.")
        for r in revisions:
            doc.add_paragraph(f"[{REVISION_LEVEL_LABELS_KO.get(RevisionLevel(r.level), r.level)}]").bold = True
            doc.add_paragraph(f"원문: {r.original_text or '-'}")
            doc.add_paragraph(f"수정안: {r.revised_text}")
            doc.add_paragraph(f"수정 사유: {r.change_reason}")
            doc.add_paragraph("")

    if is_litigation:
        doc.add_paragraph(
            "※ 본 의견서는 승소 가능성을 예측하지 않으며, 최종 대응방향은 소송대리인(변호사)의 검토를 "
            "거쳐 확정하여야 합니다."
        ).italic = True

    _write_footer_disclaimer(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_revision_request_letter(document: Document, findings: list, revisions: list) -> bytes:
    """상대방 전달용 수정 요청서 — only STANDARD-level wording is used, since this
    document leaves TOPEC's premises; internal-only STRONG wording is excluded."""
    doc = DocxDocument()
    _write_header(doc, "계약조건 수정 요청서")

    doc.add_paragraph(f"수신: {document.counterparty_name or '(상대방)'}")
    doc.add_paragraph(f"제목: {document.title} 관련 계약조건 수정 요청")
    doc.add_paragraph(
        "귀사와 체결(예정)한 위 계약과 관련하여, 아래와 같이 일부 조항에 대한 수정을 요청드립니다."
    )
    doc.add_paragraph("")

    standard_revisions = [r for r in revisions if r.level == RevisionLevel.STANDARD.value]
    if not standard_revisions:
        doc.add_paragraph("현재 권고 수정안이 없습니다.")
    for i, r in enumerate(standard_revisions, start=1):
        doc.add_heading(f"수정요청 {i}", level=2)
        doc.add_paragraph(f"현재 조항: {r.original_text or '(해당 조항 원문 미확인)'}")
        doc.add_paragraph(f"수정 요청 문구: {r.revised_text}")
        doc.add_paragraph(f"요청 사유: {r.change_reason}")
        doc.add_paragraph("")

    doc.add_paragraph("귀사의 긍정적인 검토를 부탁드립니다. 감사합니다.")
    _write_footer_disclaimer(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
